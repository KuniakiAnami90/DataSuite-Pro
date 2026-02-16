import streamlit as st
import pandas as pd
import plotly.express as px
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Sistem Analisis Data KIAS", layout="wide", page_icon="📊")

# --- CSS KHAS (UNTUK PAPARAN LEBIH KEMAS & TULISAN HITAM) ---
st.markdown("""
<style>
    /* Paksa tulisan laporan jadi hitam pekat */
    .report-view-text {
        color: black !important;
        font-family: 'Times New Roman', serif;
        background-color: white;
        padding: 30px;
        border: 1px solid #ddd;
        border-radius: 5px;
    }
    .report-view-text h1, .report-view-text h2, .report-view-text h3, 
    .report-view-text p, .report-view-text td, .report-view-text th, 
    .report-view-text li, .report-view-text span, .report-view-text div {
        color: black !important;
    }
    
    /* Table Header Styling */
    div[data-testid="stTable"] th {
        background-color: #f0f2f6 !important;
        color: black !important; 
        font-weight: bold;
    }
    
    /* Sidebar Styling */
    section[data-testid="stSidebar"] {
        background-color: #f8f9fa;
    }
</style>
""", unsafe_allow_html=True)

# --- INITIALIZATION (CLEAN SLATE - KOSONGKAN DATA) ---
if 'df' not in st.session_state:
    st.session_state['df'] = None

if 'report_structure' not in st.session_state:
    st.session_state['report_structure'] = [
        {"title": "Bab 1: Demografi Responden", "items": []}
    ]

# --- UTILITY FUNCTIONS ---

def detect_header_row(df_raw):
    """Mencari baris header sebenar berdasarkan baris yang paling banyak data."""
    max_non_na = 0
    header_idx = 0
    for i in range(min(10, len(df_raw))):
        row_count = df_raw.iloc[i].count()
        if row_count > max_non_na:
            max_non_na = row_count
            header_idx = i
    return header_idx

def clean_data(df):
    """Membersihkan data dari kolum kosong dan baris hantu."""
    # Buang kolum 'Unnamed'
    df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
    # Buang baris yang kosong sepenuhnya
    df.dropna(how='all', inplace=True)
    # Reset index
    df.reset_index(drop=True, inplace=True)
    return df

def generate_analysis_text(col_name, counts, percents):
    """Menjana ayat analisis dalam Bahasa Melayu Akademik secara automatik."""
    try:
        max_idx = counts.idxmax()
        max_val = counts[max_idx]
        max_pct = percents[max_idx]
        
        min_idx = counts.idxmin()
        min_val = counts[min_idx]
        min_pct = percents[min_idx]
        
        text = (
            f"Berdasarkan analisis deskriptif bagi pemboleh ubah **{col_name}**, dapatan kajian menunjukkan bahawa "
            f"majoriti responden memilih kategori **{max_idx}** dengan kekerapan seramai **{max_val}** orang ({max_pct:.1f}%). "
            f"Manakala, kategori **{min_idx}** mencatatkan jumlah terendah iaitu seramai **{min_val}** orang ({min_pct:.1f}%)."
        )
        return text
    except:
        return "Data tidak mencukupi untuk menjana analisis automatik."

# --- FUNGSI EXPORT WORD (DOCX) ---
def generate_word_doc(structure, df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    doc.add_heading('LAPORAN ANALISIS DATA', 0)
    
    for chapter in structure:
        doc.add_page_break()
        doc.add_heading(chapter['title'], level=1)
        
        for item in chapter['items']:
            if item['type'] == 'single':
                col = item['var']
                if col in df.columns:
                    counts = df[col].value_counts()
                    percents = df[col].value_counts(normalize=True) * 100
                    
                    # Table A (Bilangan)
                    df_A = pd.DataFrame({'Kategori': counts.index, 'Bilangan': counts.values})
                    df_A.loc[len(df_A)] = ['JUMLAH BESAR', df_A['Bilangan'].sum()]
                    
                    # Table B (Peratus)
                    df_B = pd.DataFrame({'Kategori': percents.index, 'Peratus (%)': percents.values.round(1)})
                    df_B.loc[len(df_B)] = ['JUMLAH BESAR', df_B['Peratus (%)'].sum().round(1)]
                    
                    text = generate_analysis_text(col, counts, percents)
                    
                    # Write to Doc
                    doc.add_heading(f"Analisis: {col}", level=2)
                    
                    doc.add_paragraph("Jadual (a): Taburan Kekerapan", style='Caption')
                    t1 = doc.add_table(df_A.shape[0]+1, df_A.shape[1])
                    t1.style = 'Table Grid'
                    for j, col_name in enumerate(df_A.columns):
                        t1.cell(0, j).text = col_name
                    for i, row in enumerate(df_A.itertuples(index=False)):
                        for j, val in enumerate(row):
                            t1.cell(i+1, j).text = str(val)
                    doc.add_paragraph() 

                    doc.add_paragraph("Jadual (b): Taburan Peratusan", style='Caption')
                    t2 = doc.add_table(df_B.shape[0]+1, df_B.shape[1])
                    t2.style = 'Table Grid'
                    for j, col_name in enumerate(df_B.columns):
                        t2.cell(0, j).text = col_name
                    for i, row in enumerate(df_B.itertuples(index=False)):
                        for j, val in enumerate(row):
                            t2.cell(i+1, j).text = str(val)
                    doc.add_paragraph()
                    
                    p = doc.add_paragraph(text.replace('**', ''))
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            elif item['type'] == 'cross':
                var_x, var_y = item['var_x'], item['var_y']
                if var_x in df.columns and var_y in df.columns:
                    doc.add_heading(f"Analisis Silang: {var_x} vs {var_y}", level=2)
                    
                    ct = pd.crosstab(df[var_x], df[var_y])
                    doc.add_paragraph("Jadual Silang (Crosstabulation)", style='Caption')
                    
                    # Create Table for Cross Tab
                    t3 = doc.add_table(ct.shape[0]+1, ct.shape[1]+1)
                    t3.style = 'Table Grid'
                    
                    # Headers
                    t3.cell(0, 0).text = var_x
                    for j, col_val in enumerate(ct.columns):
                        t3.cell(0, j+1).text = str(col_val)
                    
                    # Rows
                    for i, (idx_val, row) in enumerate(ct.iterrows()):
                        t3.cell(i+1, 0).text = str(idx_val)
                        for j, val in enumerate(row):
                            t3.cell(i+1, j+1).text = str(val)
                    
                    doc.add_paragraph(f"Analisis ini menunjukkan taburan silang antara {var_x} dan {var_y}.")
                    doc.add_paragraph()

    return doc

# --- SIDEBAR MENU (MENGGANTIKAN Sidebar.tsx) ---
st.sidebar.title("📊 Sistem Analisis KIAS")
menu = st.sidebar.radio("Navigasi Modul:", 
    ["1. Data Manager", "2. Dashboard", "3. Statistical Analysis", "4. Cross Analysis", "5. Report Generator", "6. Advanced Report Builder"])

# --- MODUL 1: DATA MANAGER (DataManager.tsx) ---
if menu == "1. Data Manager":
    st.title("📂 Pengurusan Data")
    st.write("Muat naik fail Excel/CSV atau Paste data anda di sini.")
    
    tab1, tab2 = st.tabs(["📤 Upload Fail", "📋 Paste (Grid View)"])
    
    with tab1:
        uploaded_file = st.file_uploader("Pilih fail (.xlsx / .csv)", type=['csv', 'xlsx'])
        if uploaded_file:
            try:
                if uploaded_file.name.endswith('.csv'):
                    df_temp = pd.read_csv(uploaded_file, header=None)
                    header_row = detect_header_row(df_temp)
                    uploaded_file.seek(0)
                    df = pd.read_csv(uploaded_file, header=header_row)
                else:
                    df = pd.read_excel(uploaded_file)
                
                df = clean_data(df)
                st.session_state['df'] = df
                st.success(f"✅ Fail berjaya dimuat naik! ({len(df)} Responden)")
                st.dataframe(df.head())
            except Exception as e:
                st.error(f"Ralat: {e}")

    with tab2:
        st.info("Klik pada sel pertama (Var1), kemudian tekan **Ctrl+V** untuk paste dari Excel.")
        if 'grid_df' not in st.session_state:
            cols = [f"Var{i+1}" for i in range(50)]
            st.session_state['grid_df'] = pd.DataFrame(columns=cols, index=range(100)).fillna("")
            
        edited_df = st.data_editor(st.session_state['grid_df'], num_rows="dynamic", use_container_width=True)
        
        if st.button("✅ Proses Data Paste"):
            try:
                clean_edit = edited_df.replace("", pd.NA).dropna(how='all', axis=0).dropna(how='all', axis=1)
                new_header = clean_edit.iloc[0]
                df_final = clean_edit[1:]
                df_final.columns = new_header
                df_final.reset_index(drop=True, inplace=True)
                st.session_state['df'] = df_final
                st.success(f"✅ Data Paste Berjaya! ({len(df_final)} Responden)")
                st.rerun()
            except Exception as e:
                st.error(f"Ralat: {e}")

# --- MODUL 2: DASHBOARD (Dashboard.tsx) ---
elif menu == "2. Dashboard":
    st.title("📈 Dashboard Ringkas")
    df = st.session_state['df']
    
    if df is None:
        st.warning("⚠️ Sila upload data dahulu di Menu 1.")
    else:
        col_list = df.columns.tolist()
        selected_col = st.selectbox("Pilih Pemboleh Ubah Utama:", col_list)
        
        c1, c2 = st.columns(2)
        with c1:
            counts = df[selected_col].value_counts()
            fig = px.bar(counts, x=counts.index, y=counts.values, 
                         labels={'x': selected_col, 'y': 'Kekerapan'}, 
                         title=f"Carta Bar: {selected_col}", color_discrete_sequence=['#3366cc'])
            st.plotly_chart(fig, use_container_width=True)
        
        with c2:
            fig2 = px.pie(counts, values=counts.values, names=counts.index, 
                          title=f"Carta Pai: {selected_col}")
            st.plotly_chart(fig2, use_container_width=True)

# --- MODUL 3: STATISTICAL ANALYSIS (StatisticalAnalysis.tsx) ---
elif menu == "3. Statistical Analysis":
    st.title("🧮 Analisis Statistik Terperinci")
    df = st.session_state['df']
    
    if df is None:
        st.warning("⚠️ Sila upload data dahulu.")
    else:
        col_list = df.columns.tolist()
        target_col = st.selectbox("Pilih Soalan / Variabel:", col_list)
        
        if target_col:
            stats = df[target_col].describe().astype(str)
            counts = df[target_col].value_counts()
            percents = df[target_col].value_counts(normalize=True) * 100
            
            c1, c2 = st.columns([1, 2])
            with c1:
                st.write("**Statistik Deskriptif:**")
                st.dataframe(stats, use_container_width=True)
            
            with c2:
                st.write("**Jadual Taburan:**")
                summary_df = pd.DataFrame({'Kategori': counts.index, 'Bilangan': counts.values, 'Peratus (%)': percents.values.round(1)})
                st.dataframe(summary_df, use_container_width=True)

# --- MODUL 4: CROSS ANALYSIS (CrossAnalysis.tsx) ---
elif menu == "4. Cross Analysis":
    st.title("❌ Analisis Silang (Crosstab)")
    df = st.session_state['df']
    
    if df is None:
        st.warning("⚠️ Sila upload data dahulu.")
    else:
        col_list = df.columns.tolist()
        c1, c2 = st.columns(2)
        with c1: x_var = st.selectbox("Paksi X (Faktor):", col_list, index=0)
        with c2: y_var = st.selectbox("Paksi Y (Hasil):", col_list, index=min(1, len(col_list)-1))
        
        if x_var and y_var:
            ct = pd.crosstab(df[x_var], df[y_var])
            
            st.subheader("Peta Haba (Heatmap)")
            fig = px.imshow(ct, text_auto=True, aspect="auto", color_continuous_scale="Blues", title=f"Hubungan {x_var} vs {y_var}")
            st.plotly_chart(fig, use_container_width=True)
            
            st.subheader("Jadual Data")
            st.dataframe(ct)

# --- MODUL 5: REPORT GENERATOR (ReportGenerator.tsx) ---
elif menu == "5. Report Generator":
    st.title("📑 Auto Report Generator")
    df = st.session_state['df']
    
    if df is None:
        st.warning("⚠️ Tiada data.")
    else:
        st.info("Laporan ini dijana secara automatik untuk SEMUA pemboleh ubah.")
        cols_to_analyze = [c for c in df.columns if c.lower() not in ['timestamp', 'id', 'email']]
        
        st.markdown('<div class="report-view-text">', unsafe_allow_html=True)
        st.header("LAPORAN ANALISIS PENUH")
        
        for col in cols_to_analyze:
            st.subheader(f"Analisis: {col}")
            counts = df[col].value_counts()
            percents = df[col].value_counts(normalize=True) * 100
            
            df_A = pd.DataFrame({'Kategori': counts.index, 'Bilangan': counts.values})
            df_A.loc[len(df_A)] = ['JUMLAH BESAR', df_A['Bilangan'].sum()]
            
            df_B = pd.DataFrame({'Kategori': percents.index, 'Peratus (%)': percents.values.round(1)})
            df_B.loc[len(df_B)] = ['JUMLAH BESAR', df_B['Peratus (%)'].sum().round(1)]

            c1, c2 = st.columns(2)
            with c1: 
                st.write("**Jadual (a): Kekerapan**")
                st.table(df_A)
            with c2: 
                st.write("**Jadual (b): Peratusan**")
                st.table(df_B)
            
            text = generate_analysis_text(col, counts, percents)
            st.markdown(f"<p style='text-align: justify;'>{text}</p>", unsafe_allow_html=True)
            st.markdown("---")
        st.markdown('</div>', unsafe_allow_html=True)

# --- MODUL 6: ADVANCED REPORT (AdvancedReportGenerator.tsx) ---
elif menu == "6. Advanced Report Builder":
    st.title("📝 Advanced Report Builder (Bab demi Bab)")
    df = st.session_state['df']
    
    if df is None:
        st.warning("⚠️ Sila upload data dahulu.")
    else:
        # Konfigurasi Bab
        with st.sidebar:
            st.divider()
            if st.button("➕ Tambah Bab Baru"):
                new_num = len(st.session_state['report_structure']) + 1
                st.session_state['report_structure'].append({"title": f"Bab {new_num}: (Klik untuk Edit)", "items": []})
                st.rerun()
            if st.button("🗑️ Reset Semua Bab"):
                st.session_state['report_structure'] = [{"title": "Bab 1: Pendahuluan", "items": []}]
                st.rerun()

        # Builder UI
        for i, chapter in enumerate(st.session_state['report_structure']):
            with st.expander(f"{chapter['title']}", expanded=True):
                new_title = st.text_input(f"Tajuk Bab {i+1}", value=chapter['title'], key=f"title_{i}")
                st.session_state['report_structure'][i]['title'] = new_title
                
                # List Items
                if chapter['items']:
                    for j, item in enumerate(chapter['items']):
                        if item['type'] == 'single':
                            st.text(f"{j+1}. Single Variable: {item['var']}")
                        else:
                            st.text(f"{j+1}. Cross Analysis: {item['var_x']} VS {item['var_y']}")
                else:
                    st.caption("Tiada item lagi.")

                # Add Item UI
                st.markdown("---")
                c1, c2, c3 = st.columns([1, 2, 1])
                with c1: type_choice = st.selectbox("Jenis Analisis", ["Single Variable", "Cross Analysis"], key=f"type_{i}")
                with c2:
                    col_opts = df.columns.tolist()
                    if type_choice == "Single Variable":
                        var_sel = st.selectbox("Pilih Variable", col_opts, key=f"v1_{i}")
                    else:
                        var_x = st.selectbox("Paksi X", col_opts, key=f"vx_{i}")
                        var_y = st.selectbox("Paksi Y", col_opts, key=f"vy_{i}")
                with c3:
                    st.write("")
                    st.write("")
                    if st.button("➕ Tambah Item", key=f"add_{i}"):
                        item = {"type": "single", "var": var_sel} if type_choice == "Single Variable" else {"type": "cross", "var_x": var_x, "var_y": var_y}
                        chapter['items'].append(item)
                        st.rerun()

        # Export Section
        st.divider()
        st.subheader("Muat Turun Laporan")
        if st.button("📄 Generate Report Preview"):
            st.markdown('<div class="report-view-text">', unsafe_allow_html=True)
            for chapter in st.session_state['report_structure']:
                st.header(chapter['title'])
                for item in chapter['items']:
                    if item['type'] == 'single':
                        col = item['var']
                        if col in df.columns:
                            st.subheader(f"Analisis: {col}")
                            counts = df[col].value_counts()
                            percents = df[col].value_counts(normalize=True) * 100
                            
                            df_A = pd.DataFrame({'Kategori': counts.index, 'Bilangan': counts.values})
                            df_A.loc[len(df_A)] = ['JUMLAH BESAR', df_A['Bilangan'].sum()]
                            
                            df_B = pd.DataFrame({'Kategori': percents.index, 'Peratus (%)': percents.values.round(1)})
                            df_B.loc[len(df_B)] = ['JUMLAH BESAR', df_B['Peratus (%)'].sum().round(1)]
                            
                            c1, c2 = st.columns(2)
                            with c1: st.table(df_A)
                            with c2: st.table(df_B)
                            
                            text = generate_analysis_text(col, counts, percents)
                            st.write(text)
            st.markdown('</div>', unsafe_allow_html=True)

        # Word Download Button
        doc = generate_word_doc(st.session_state['report_structure'], df)
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label="📝 Muat Turun Laporan (.docx)",
            data=bio.getvalue(),
            file_name="Laporan_Analisis_KIAS.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
